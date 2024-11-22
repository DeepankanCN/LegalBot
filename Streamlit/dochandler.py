import os
import time
import asyncio
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import ChatOpenAI
from langchain_text_splitters import CharacterTextSplitter
from langchain.chains import LLMChain
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document 
from alj_analysis import format_time, markdown_to_plain_text
from docx.shared import Pt
from langchain_core.documents import Document as dt
from typing import List, Dict, Set
from collections import OrderedDict
import logging
import fitz

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv(override=True)
api_key = os.getenv('open_ai_key')

class DocumentProcessor:
    def __init__(self, max_retries=3, retry_delay=1):
        self.max_retries = max_retries
        self.retry_delay = retry_delay
        self.pending_docs = asyncio.Queue()
        self.results = OrderedDict()
        self.processing = set()
        self.failed_docs = set()
        self.completed = 0
        self.total_docs = 0
        self.running = True
        self.active_tasks = set()

    async def process_single_doc(self, chain: LLMChain, doc: str, index: int, attempt: int = 1) -> tuple[int, str, bool]:
        """Process a single document with retry logic"""
        try:
            result = await chain.arun(document_text=doc)
            logger.info(f"Successfully processed document {index} on attempt {attempt}")
            return index, result, True
        except Exception as e:
            logger.error(f"Error processing document {index} (attempt {attempt}): {str(e)}")
            if attempt < self.max_retries:
                logger.info(f"Retrying document {index} in {self.retry_delay} seconds...")
                await asyncio.sleep(self.retry_delay)
                return await self.process_single_doc(chain, doc, index, attempt + 1)
            return index, f"Failed to process after {self.max_retries} attempts: {str(e)}", False

    async def process_documents(self, chain: LLMChain):
        """Continuously process documents from the queue"""
        while self.running or not self.pending_docs.empty():
            try:
                doc_data = await self.pending_docs.get()
                if doc_data is None:
                    continue

                index, doc = doc_data
                if index in self.processing:
                    continue

                self.processing.add(index)
                index, result, success = await self.process_single_doc(chain, doc, index)
                
                if success:
                    self.results[index] = result
                    self.completed += 1
                    logger.info(f"Completed {self.completed}/{self.total_docs} documents")
                else:
                    self.failed_docs.add((index, doc))
                
                self.processing.remove(index)
                self.pending_docs.task_done()

            except Exception as e:
                logger.error(f"Error in process_documents: {str(e)}")
                continue

    async def process_failed_documents(self, chain: LLMChain):
        """Retry failed documents"""
        while self.failed_docs:
            index, doc = self.failed_docs.pop()
            await self.pending_docs.put((index, doc))

def chunk_control(chunk_size, pdf_name):
    """Control the chunking of PDF documents"""
    def extract_text_from_pdf(pdf_path):
        text_list = []
        with fitz.open(pdf_path) as pdf:
            for page_num in range(len(pdf)):
                page = pdf.load_page(page_num)
                text = page.get_text()
                text_list.append(text)
        return text_list

    extracted_text = extract_text_from_pdf(pdf_name)
    print(len(extracted_text))

    text_splitter = CharacterTextSplitter(
        separator="\n", chunk_size=chunk_size, chunk_overlap=4000, length_function=len)
    result = '\n'.join(doc for doc in extracted_text)
    docs = text_splitter.split_text(result)
    print(len(docs))
    return docs, extracted_text

def new_chunk_control(docs, chunk_size):
    """Control chunking for already extracted text"""
    text_splitter = CharacterTextSplitter(
        separator="\n", chunk_size=chunk_size, chunk_overlap=4000, length_function=len)
    
    print(len(docs))
    result = '\n'.join(doc for doc in docs)
    docs = text_splitter.split_text(result)
    print(len(docs))
    return docs

def create_word_doc(strings, folder='.', file_name='section2.docx'):
    """Create a Word document from the processed strings"""
    if not os.path.exists(folder):
        os.makedirs(folder)

    file_path = os.path.join(folder, file_name)
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)

    for text in strings:
        doc.add_paragraph(text)
        if len(doc.paragraphs) > 0:
            last_paragraph = doc.paragraphs[-1]
            if len(last_paragraph.text) > 2000:
                doc.add_page_break()

    doc.save(file_path)
    print(f"Document saved as {file_path}")

async def structureformat_async(docs, filename, titlelist):
    """Async implementation of document structuring"""
    start_time = time.time()
    concurrent_tasks = 5
    
    llm = ChatOpenAI(
        openai_api_key=api_key,
        model="gpt-4o-mini",
        temperature=0,
        max_retries=2,
        streaming=True,
        verbose=True
    )

    prompt = """
Extract the following from the unstructured data:
- Exhibit Title
- Page Number
- Content (Get the important cotent in readable format,This is for legal purpose)
Unstructured Data:
{document_text}

Output:
Exhibit Title: Exhibit X,
 Page Number: Page Y, 
 Content: Content text here." 
"""

    prompt_template = PromptTemplate(
        input_variables=["document_text"],
        template=prompt
    )
    chain = LLMChain(llm=llm, prompt=prompt_template, verbose=False)

    # Initialize document processor
    processor = DocumentProcessor(max_retries=3, retry_delay=1)
    processor.total_docs = len(docs)

    # Add documents to queue
    for i, doc in enumerate(docs):
        await processor.pending_docs.put((i, doc))

    # Create processing tasks
    tasks = []
    for _ in range(concurrent_tasks):
        task = asyncio.create_task(processor.process_documents(chain))
        tasks.append(task)

    # Wait for initial processing
    await processor.pending_docs.join()

    # Process failed documents
    while processor.failed_docs:
        await processor.process_failed_documents(chain)
        await processor.pending_docs.join()

    # Clean up
    processor.running = False
    for task in tasks:
        task.cancel()
    await asyncio.gather(*tasks, return_exceptions=True)

    # Process results
    textlist = [processor.results[i] for i in range(len(docs)) if i in processor.results]
    new_textlist = []
    for text in textlist:
        k = markdown_to_plain_text(text)
        new_textlist.append(k)

    # Create output document
    create_word_doc(new_textlist, folder="StructuredOutput", file_name=filename)

    elapsed_time = time.time() - start_time
    logger.info(f"Processing completed in {elapsed_time:.2f} seconds")
    logger.info(f"Successfully processed {len(processor.results)}/{len(docs)} documents")
    
    if len(processor.results) < len(docs):
        logger.warning(f"Failed to process {len(docs) - len(processor.results)} documents")

    return 1

# def structureformat(docs, filename, titlelist):
#     """Main entry point for document structuring"""
#     return asyncio.run(structureformat_async(docs, filename, titlelist))



# def structureformat(docs, filename, titlelist):
#     """Main entry point for document structuring"""
#     try:
#         # Run the async function using asyncio.run
#         result = asyncio.run(structureformat_async(docs, filename, titlelist))
#         return result
#     except Exception as e:
#         logger.error(f"Error in structureformat: {str(e)}")
#         raise


async def structureformat(docs, filename, titlelist):
    """Main entry point for document structuring"""
    try:
        result = await structureformat_async(docs, filename, titlelist)
        return result
    except Exception as e:
        logger.error(f"Error in structureformat: {str(e)}")
        raise

