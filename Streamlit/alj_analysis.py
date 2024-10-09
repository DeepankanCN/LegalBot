import os
import time
from langchain_community.document_loaders import PyPDFLoader
from langchain.chains.summarize import load_summarize_chain
from langchain_openai import ChatOpenAI
from langchain_core.documents import Document
from langchain_text_splitters import CharacterTextSplitter
from langchain.chains import RefineDocumentsChain, LLMChain
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Pt

load_dotenv(override=True)

api_key = os.getenv('open_ai_key')
llm = ChatOpenAI(openai_api_key=api_key, model="gpt-4o-mini", temperature=0)

def analyze_alj_decision(file_path):
    start_time = time.time()
    
    # Load the file
    loader = PyPDFLoader(file_path)
    docs = loader.load()
    
    # Initialize text splitter
    text_splitter = CharacterTextSplitter(separator="\n", chunk_size=100000, chunk_overlap=100)
    result = '\n'.join(doc.page_content for doc in docs)
    docs = text_splitter.split_text(result)

    # Create Langchain Document Objects
    temp_doc = [Document(page_content=doc) for doc in docs]
    docs = temp_doc

    # Define prompts
    document_prompt = PromptTemplate(input_variables=["page_content"], template="{page_content}")
    
    prompt = PromptTemplate.from_template(
        "Based on the following content, Try to find loopholes in judge decision that made it unfavourable, I want some flaws in the decision so that I can get favourable decision later on, Dont write this in Markdown : {context}"
    "Make it as detailed as possible"
    )
    initial_llm_chain = LLMChain(llm=llm, prompt=prompt)
    initial_response_name = "prev_response"

    # Updated prompt for refinement to enhance and finalize the questions, ensuring a total of 20 questions
    prompt_refine = PromptTemplate.from_template(
        "Here are the intial info from the document: {prev_response}. "
        "Refine these questions and ensure the final set contains exactly questions based on the following additional context: {context}"
        "Quote each and everything, find the smallest details"
        "Dont write this in markdown format"
    )
    refine_llm_chain = LLMChain(llm=llm, prompt=prompt_refine)

    chain = RefineDocumentsChain(
        initial_llm_chain=initial_llm_chain,
        refine_llm_chain=refine_llm_chain,
        document_prompt=document_prompt,
        document_variable_name="context",
        initial_response_name="prev_response",
    )

    # Run the analysis
    k = chain.run(docs)

    # Calculate elapsed time
    elapsed_time = time.time() - start_time
    formatted_time = format_time(elapsed_time)

    # Create Word document
    plain_text = markdown_to_plain_text(k)
    create_word_doc(plain_text, "loopholes.docx")
    
    return formatted_time, "Loopholes found"

def format_time(seconds):
    """Format time in minutes and seconds."""
    minutes, seconds = divmod(int(seconds), 60)
    return f"{minutes} min {seconds} sec"

def markdown_to_plain_text(markdown_text):
    """Convert Markdown text to plain text."""
    html = markdown(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')
    return soup.get_text()

def create_word_doc(text, file_name='loopholes.docx'):
    """Create a Word document with the provided plain text."""
    doc = DocxDocument()
    
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)

    doc.add_paragraph(text)
    doc.save(file_name)
    print(f"Document saved as {file_name}")
