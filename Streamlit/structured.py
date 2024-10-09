
import os
import time
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import ChatOpenAI
from langchain_text_splitters import CharacterTextSplitter
from langchain.chains import LLMChain
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document 
from alj_analysis import format_time,markdown_to_plain_text
from docx.shared import Pt

load_dotenv(override=True)

api_key = os.getenv('open_ai_key')
llm = ChatOpenAI(openai_api_key=api_key, model="gpt-4o-mini", temperature=0)




def chunk_control(chunk_size,pdf_name):
    import fitz  # PyMuPDF

    def extract_text_from_pdf(pdf_path):
        # List to store the text of each page
        text_list = []

        # Open the PDF file
        with fitz.open(pdf_path) as pdf:
            # Iterate over each page
            for page_num in range(len(pdf)):
                page = pdf.load_page(page_num)
                # Extract text from the page
                text = page.get_text()
                # Append text to the list
                text_list.append(text)

        return text_list

    # Example usage
    #pdf_name = 'section_5.pdf'  # Replace with your PDF file name
    extracted_text = extract_text_from_pdf(pdf_name)


    print(len(extracted_text))

    


    #Initiating a splitter function
    text_splitter = CharacterTextSplitter(
        separator="\n", chunk_size=chunk_size, chunk_overlap=4000, length_function=len)
    result=''
    result = '\n'.join(doc for doc in extracted_text)
    docs = text_splitter.split_text(result)
    print(len(docs))
    return docs,extracted_text




def structureformat(docs,filename):
    start_time = time.time()

    
    textlist=[]
    print("Doc len is :", len(docs))
    for i in range(len(docs)):
        

    #Define the prompt
        prompt = """
        Transform the following unstructured data into a well-organized and content-rich textual format. Extract all key information and arrange it logically using headings, subheadings, and bullet points as needed. Ensure that all details from the original text are preserved and presented clearly.

        Unstructured Data:

        {document_text}

        Requirements:

        Headings and Subheadings:

        Use headings and subheadings to categorize different sections of the information.
        Bullet Points and Lists:

        Utilize bullet points or numbered lists to break down complex information or lists of items.
        Clarity and Coherence:

        Ensure the final text is easy to read and understand while retaining all original details.
        Accuracy:

        Maintain the accuracy and context of the information from the original data.




        Output:
        """



        # Define the prompt template
        prompt_template = PromptTemplate(
            input_variables=["document_text"],
            template=prompt
        )
        text=docs[i]

        # Create the chain
        chain = LLMChain(llm=llm, prompt=prompt_template)

        # Run the chain with your document text
        result = chain.run(document_text=text)
        print(i)

        # Output the result
        textlist.append(result)

    new_textlist=[]
    for i in textlist:
        k=''
        k=markdown_to_plain_text(i)
        new_textlist.append(k)
    create_word_doc(new_textlist, folder="StructuredOutput", file_name=filename)
    

        # Calculate elapsed time
    elapsed_time = time.time() - start_time
    formatted_time = format_time(elapsed_time)
    return formatted_time




def create_word_doc(strings, folder='.', file_name='section2.docx'):
    # Ensure the folder exists
    if not os.path.exists(folder):
        os.makedirs(folder)

    # Create the full file path
    file_path = os.path.join(folder, file_name)

    doc = Document()
    
    # Set default font size for the document
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)

    # Add content to the document
    for text in strings:
        doc.add_paragraph(text)

        # Simple pagination check
        if len(doc.paragraphs) > 0:
            last_paragraph = doc.paragraphs[-1]
            if len(last_paragraph.text) > 2000:  # Arbitrary number, tune as needed
                doc.add_page_break()

    # Save the document
    doc.save(file_path)
    print(f"Document saved as {file_path}")






